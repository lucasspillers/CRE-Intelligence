from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1A, 0x23, 0x32)
GRAY = RGBColor(0x4B, 0x55, 0x63)
MUTED = RGBColor(0x9C, 0xA3, 0xAF)
CARD_FILL = "F7F8FA"


def plotly_fig_to_image_bytes(fig):
    return fig.to_image(format="png", width=900, height=500, scale=2)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_divider(doc):
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(18)


def add_body_text(doc, text):
    for paragraph_text in text.split("\n"):
        if not paragraph_text.strip():
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_text.strip())
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY


def add_highlight_callouts(doc, highlights):
    for bullet in highlights:
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade_cell(cell, CARD_FILL)
        run = cell.paragraphs[0].add_run(bullet.strip())
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)


def add_metric_grid(doc, metric_pairs, columns=3):
    rows_needed = -(-len(metric_pairs) // columns)
    table = doc.add_table(rows=rows_needed, cols=columns)
    idx = 0
    for r in range(rows_needed):
        for c in range(columns):
            if idx >= len(metric_pairs):
                break
            label, value = metric_pairs[idx]
            cell = table.rows[r].cells[c]
            _shade_cell(cell, CARD_FILL)
            p1 = cell.paragraphs[0]
            run1 = p1.add_run(label.upper())
            run1.font.size = Pt(8)
            run1.font.color.rgb = MUTED
            p2 = cell.add_paragraph()
            run2 = p2.add_run(value)
            run2.font.size = Pt(14)
            run2.font.bold = True
            run2.font.color.rgb = NAVY
            idx += 1
    doc.add_paragraph()


def add_property_specifications(doc, property_data):
    prop = property_data["property"]
    spec_pairs = []
    if prop.get("square_feet"):
        spec_pairs.append(("Square Feet", f"{prop['square_feet']:,.0f} SF"))
    if prop.get("year_built"):
        spec_pairs.append(("Year Built", str(int(prop["year_built"]))))
    if prop.get("occupancy_pct") is not None:
        spec_pairs.append(("Occupancy", f"{prop['occupancy_pct']}%"))
    if prop.get("vacancy_rate_pct") is not None:
        spec_pairs.append(("Vacancy", f"{prop['vacancy_rate_pct']}%"))
    if prop.get("type"):
        spec_pairs.append(("Property Type", prop["type"]))

    if spec_pairs:
        add_metric_grid(doc, spec_pairs, columns=3)
    else:
        add_body_text(doc, "Detailed property specifications have not yet been provided.")


def add_property_photos(doc, photos):
    if not photos:
        return
    for photo_bytes in photos:
        try:
            doc.add_picture(BytesIO(photo_bytes), width=Inches(6.5))
            doc.add_paragraph()
        except Exception:
            continue


def build_offering_memorandum(property_data, om_narrative, bar_fig, line_fig, property_photos=None):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Cover Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(property_data["property"]["name"] or "Untitled Property")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Offering Memorandum").font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = GRAY

    address_line = doc.add_paragraph()
    address_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    address_text = f"{property_data['property']['address']}, {property_data['property']['city']}, {property_data['property']['state']}"
    address_line.add_run(address_text).font.color.rgb = GRAY

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_line.add_run(f"Prepared {date.today().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = MUTED

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = confidential.add_run("Confidential — Prepared for Investment Analysis Purposes Only")
    conf_run.font.size = Pt(9)
    conf_run.italic = True

    if property_photos:
        doc.add_paragraph()
        add_property_photos(doc, property_photos)

    doc.add_page_break()

    # The Story
    add_section_heading(doc, "Executive Summary")
    add_body_text(doc, om_narrative.get("executive_summary", ""))
    add_divider(doc)

    add_section_heading(doc, "Property Overview")
    add_body_text(doc, om_narrative.get("property_overview", ""))
    add_divider(doc)

    add_section_heading(doc, "Property Specifications")
    add_property_specifications(doc, property_data)
    add_divider(doc)

    add_section_heading(doc, "Market Overview")
    add_body_text(doc, om_narrative.get("market_overview", ""))
    add_divider(doc)

    add_section_heading(doc, "Tenant Overview")
    add_body_text(doc, om_narrative.get("tenant_overview", ""))
    add_divider(doc)

    add_section_heading(doc, "Investment Highlights")
    add_highlight_callouts(doc, om_narrative.get("investment_highlights", []))

    doc.add_page_break()

        

    # The Numbers
    add_section_heading(doc, "Financial Summary")
    metrics = property_data["metrics"]
    financing = property_data["financing"]
    metric_pairs = [
        ("Purchase Price", f"${financing['purchase_price']:,.0f}"),
        ("Loan Amount", f"${financing['loan_amount']:,.0f}"),
        ("Initial Equity", f"${financing['initial_equity']:,.0f}"),
        ("Purchase Cap Rate", f"{metrics['purchase_cap_rate_pct']}%"),
        ("Cash-on-Cash Return (Yr 1)", f"{metrics['cash_on_cash_return_pct']}%"),
        ("Equity Multiple", f"{metrics['equity_multiple']}x"),
        ("Debt Yield", f"{metrics['debt_yield_pct']}%"),
        ("DSCR", f"{metrics['dscr']}x"),
        ("Estimated Exit Value", f"${metrics['estimated_exit_value']:,.0f}"),
    ]
    add_metric_grid(doc, metric_pairs, columns=3)
    add_divider(doc)

    add_section_heading(doc, "Cash Flow Schedule")
    cf_table = doc.add_table(rows=1, cols=4)
    cf_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Year", "NOI", "Debt Service", "Cash Flow"]):
        cf_table.rows[0].cells[i].text = h
    for row in property_data["cash_flow_schedule"]:
        row_cells = cf_table.add_row().cells
        row_cells[0].text = str(row["Year"])
        row_cells[1].text = f"${row['NOI']:,.0f}"
        row_cells[2].text = f"${row['Debt Service']:,.0f}"
        row_cells[3].text = f"${row['Cash Flow']:,.0f}"

    doc.add_paragraph()
    add_body_text(doc, om_narrative.get("cash_flow_commentary", ""))
    doc.add_page_break()

    add_section_heading(doc, "Financial Performance")
    doc.add_picture(BytesIO(plotly_fig_to_image_bytes(bar_fig)), width=Inches(6.5))
    doc.add_paragraph()
    doc.add_picture(BytesIO(plotly_fig_to_image_bytes(line_fig)), width=Inches(6.5))
    doc.add_page_break()

    add_section_heading(doc, "Disclaimer")
    disclaimer_text = (
        "This Offering Memorandum has been prepared for informational purposes only and does not "
        "constitute an offer to sell or a solicitation of an offer to buy any security or investment. "
        "The information contained herein has been obtained from sources believed to be reliable, but "
        "no representation or warranty, express or implied, is made as to the accuracy or completeness "
        "of such information. All financial projections are based on assumptions and estimates that are "
        "subject to significant uncertainty and should not be relied upon as a guarantee of future "
        "performance. Prospective investors should conduct their own independent due diligence, including "
        "engagement of qualified legal, tax, and financial advisors, prior to making any investment decision."
    )
    add_body_text(doc, disclaimer_text)

    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

